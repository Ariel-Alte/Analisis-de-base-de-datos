"""
App Streamlit — Análisis de Informes de Mantenimiento
======================================================
Instalación (una sola vez):
    pip install streamlit plotly openpyxl pandas numpy python-docx

Ejecución:
    streamlit run app_mantenimiento.py
"""

import io
import re
import json
import numpy as np
import pandas as pd
import streamlit as st
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# ─────────────────────────────────────────────
# CONFIGURACIÓN DE PÁGINA
# ─────────────────────────────────────────────
st.set_page_config(
    page_title="Análisis de Mantenimiento",
    page_icon="🔧",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ─────────────────────────────────────────────
# ESTILOS CSS
# ─────────────────────────────────────────────
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Sans:wght@400;600;700&family=IBM+Plex+Mono:wght@400;600&display=swap');

    html, body, [class*="css"] { font-family: 'IBM Plex Sans', sans-serif; }

    .stApp { background-color: #0f1117; color: #e8eaf0; }

    /* Sidebar */
    section[data-testid="stSidebar"] {
        background: #161b27;
        border-right: 1px solid #2a3040;
    }

    /* Títulos */
    h1 { color: #4fc3f7 !important; font-weight: 700 !important; letter-spacing: -0.5px; }
    h2 { color: #81d4fa !important; font-weight: 600 !important; }
    h3 { color: #b0bec5 !important; font-weight: 600 !important; }

    /* KPI Cards */
    .kpi-card {
        background: linear-gradient(135deg, #1a2235 0%, #1e2a3a 100%);
        border: 1px solid #2a3a50;
        border-left: 4px solid #4fc3f7;
        border-radius: 8px;
        padding: 20px 24px;
        text-align: center;
        transition: border-color 0.2s;
    }
    .kpi-card:hover { border-left-color: #81d4fa; }
    .kpi-value {
        font-family: 'IBM Plex Mono', monospace;
        font-size: 2.4rem;
        font-weight: 700;
        color: #4fc3f7;
        line-height: 1;
        margin-bottom: 6px;
    }
    .kpi-label {
        font-size: 0.78rem;
        color: #78909c;
        text-transform: uppercase;
        letter-spacing: 0.08em;
    }
    .kpi-card.danger  { border-left-color: #ef5350; }
    .kpi-card.danger .kpi-value { color: #ef5350; }
    .kpi-card.warning { border-left-color: #ffa726; }
    .kpi-card.warning .kpi-value { color: #ffa726; }
    .kpi-card.success { border-left-color: #66bb6a; }
    .kpi-card.success .kpi-value { color: #66bb6a; }

    /* Section divider */
    .section-header {
        border-bottom: 2px solid #2a3a50;
        padding-bottom: 8px;
        margin: 32px 0 16px 0;
        color: #4fc3f7;
        font-size: 1.1rem;
        font-weight: 700;
        text-transform: uppercase;
        letter-spacing: 0.06em;
    }

    /* Desviación badge */
    .desv-pos { color: #ef5350; font-weight: 700; font-family: 'IBM Plex Mono', monospace; }
    .desv-neg { color: #42a5f5; font-weight: 700; font-family: 'IBM Plex Mono', monospace; }
    .crit-c   { background: #b71c1c; color: white; padding: 2px 8px; border-radius: 4px; font-size: 0.8rem; font-weight: 700; }
    .crit-r   { background: #e65100; color: white; padding: 2px 8px; border-radius: 4px; font-size: 0.8rem; font-weight: 700; }
    .crit-n   { background: #263238; color: #90a4ae; padding: 2px 8px; border-radius: 4px; font-size: 0.8rem; }

    /* Upload area */
    .uploadedFile { background: #1a2235 !important; border: 1px solid #2a3a50 !important; }

    /* Dataframe */
    .stDataFrame { border-radius: 8px; overflow: hidden; }

    /* Tabs */
    .stTabs [data-baseweb="tab-list"] { gap: 8px; background: transparent; }
    .stTabs [data-baseweb="tab"] {
        background: #1a2235;
        border-radius: 6px 6px 0 0;
        color: #78909c;
        padding: 8px 20px;
        border: 1px solid #2a3a50;
        border-bottom: none;
    }
    .stTabs [aria-selected="true"] {
        background: #1e2a3a !important;
        color: #4fc3f7 !important;
        border-color: #4fc3f7 !important;
    }

    /* Selectbox, multiselect */
    .stSelectbox > div, .stMultiSelect > div { background: #1a2235 !important; }

    /* Info/warning boxes */
    .stAlert { border-radius: 8px; }

    /* Scrollbar */
    ::-webkit-scrollbar { width: 6px; height: 6px; }
    ::-webkit-scrollbar-track { background: #0f1117; }
    ::-webkit-scrollbar-thumb { background: #2a3a50; border-radius: 3px; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────
# FUNCIONES DE ANÁLISIS
# ─────────────────────────────────────────────

MONTH_ORDER = [
    'ENERO','FEBRERO','MARZO','ABRIL','MAYO','JUNIO',
    'JULIO','AGOSTO','SEPTIEMBRE','OCTUBRE','NOVIEMBRE','DICIEMBRE'
]

SISTEMA_LABELS = {
    'BG' : 'Bogie',
    'SLN': 'Salón',
    'EXT': 'Exterior',
    'TYC': 'Tracción y Choque',
    'PM' : 'Par Montado',
    'EBC': 'Elementos bajo coche',
    'NSF': 'Sistema de freno Neumatico',
    'MSF': 'Sistema de freno Mecanico',
    'SFM': 'Sistema de freno Mecanico',
    'DSM': 'Sala de motor Diesel',
    'CAB': 'Cabina',
    'ATS': 'ATS',
    'DOC': 'Documentación',
    'NGN': 'Ninguna observación en el informe',
}

def parse_multivalor(v):
    """
    Extrae TODOS los valores numéricos de una celda.
    Soporta separadores / y _ además de comas como decimal.
    Ejemplos:
        '238/241'        → [238.0, 241.0]
        '1,5/3,5/3,5'   → [1.5, 3.5, 3.5]
        '31/28/32/30'   → [31.0, 28.0, 32.0, 30.0]
        427              → [427.0]
        None             → []
    """
    if v is None:
        return []
    if isinstance(v, (int, float)):
        return [float(v)] if not np.isnan(float(v)) else []
    s = str(v).strip()
    if not s:
        return []
    # Normalizar: reemplazar coma decimal por punto, luego separar por / o _
    # Distinguir coma decimal (ej: 1,5) de coma separador (ej: 1,5/3,5)
    # Estrategia: separar primero por / y _, luego limpiar cada token
    tokens = re.split(r'[/_]', s)
    resultado = []
    for tok in tokens:
        tok = tok.strip().replace(',', '.')
        nums = re.findall(r'-?[\d]+\.?[\d]*', tok)
        for n in nums:
            try:
                resultado.append(float(n))
            except ValueError:
                pass
    return resultado

def parse_valor(v):
    """Compatibilidad: retorna el primer valor numérico (para usos simples)."""
    vals = parse_multivalor(v)
    return vals[0] if vals else np.nan

def calcular_desvio(row):
    """
    Evalúa TODOS los valores de R1 y R2 (separados por / o _)
    y retorna el par (valor_relevado, desvio) con el MAYOR desvío absoluto.
    Así el informe siempre reporta el peor caso de todos los puntos medidos.
    """
    # Extraer todos los valores de ambas celdas
    todos = parse_multivalor(row.get('Relevado1')) + parse_multivalor(row.get('Relevado2'))
    if not todos:
        return np.nan, np.nan

    ref_min = row['RefMin_num'] if pd.notna(row['RefMin_num']) else None
    ref_max = row['RefMax_num'] if pd.notna(row['RefMax_num']) else None

    mejor_dev, mejor_val = 0.0, todos[0]

    for v in todos:
        # Desvío por debajo del mínimo (negativo)
        if ref_min is not None and v < ref_min:
            dev = v - ref_min
            if abs(dev) > abs(mejor_dev):
                mejor_dev, mejor_val = dev, v
        # Desvío por encima del máximo (positivo)
        if ref_max is not None and v > ref_max:
            dev = v - ref_max
            if abs(dev) > abs(mejor_dev):
                mejor_dev, mejor_val = dev, v

    return mejor_val, mejor_dev

@st.cache_data(show_spinner=False)
def cargar_y_analizar(file_bytes):
    """
    Lee el Excel y devuelve (df, df_out).
    Auto-detecta dos formatos:
      Formato A: dos bloques en paralelo, con columnas RefMin/RefMax/Relevado
      Formato B: un solo bloque, sin columnas de referencia (ej: Roca)
    """
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active

    # ── 1. Detectar fila de headers ──
    header_row = 2
    for i, row in enumerate(ws.iter_rows(min_row=1, max_row=10, values_only=True), start=1):
        if row[1] is not None:
            header_row = i
            break
    data_start = header_row + 1

    # ── 2. Leer headers reales ──
    headers_raw = [ws.cell(row=header_row, column=c).value
                   for c in range(1, ws.max_column + 1)]

    # ── 3. Detectar si hay segundo bloque (se repite el primer header) ──
    first_header = next((h for h in headers_raw if h is not None), None)
    b1_start = next((i for i, h in enumerate(headers_raw) if h is not None), 1)
    bloque2_col = None
    if first_header:
        found = False
        for i, h in enumerate(headers_raw):
            if h == first_header:
                if not found:
                    found = True
                else:
                    bloque2_col = i
                    break

    b1_end = bloque2_col if bloque2_col else len(headers_raw)

    # ── 4. Detectar formato segun headers del bloque 1 ──
    hdrs_b1 = [h for h in headers_raw[b1_start:b1_end] if h is not None]
    tiene_refs = any('Referencia' in str(h) or 'Relevado' in str(h) for h in hdrs_b1)

    # ── 5. Leer filas ──
    rows1, rows2 = [], []
    for row in ws.iter_rows(min_row=data_start, values_only=True):
        r1 = row[b1_start:b1_end]
        if any(v is not None for v in r1):
            rows1.append(r1)
        if bloque2_col is not None:
            b2_start = bloque2_col
            b2_end = b2_start + (b1_end - b1_start)
            if b2_end <= len(row):
                r2 = row[b2_start:b2_end]
                if any(v is not None for v in r2):
                    rows2.append(r2)

    # ── 6. Definir columnas segun formato ──
    if tiene_refs:
        # Formato A: con RefMin, RefMax, Relevado1, Relevado2
        cols = [
            'Mes','Responsable','Contrato','Linea','Vehiculo','Modulo','MR',
            'Modelo','Servicio','Fecha','NroInforme','SistemaUnidad',
            'SistemaAmpliado','Item1','Item2','Descripcion',
            'RefMin','RefMax','Relevado1','Relevado2','Criticidad',
            'DescAgrupada','CritAmpliado','CodItem','FechaReInsp','NroReInsp',
            'SistUnitReInsp','SistAmpReInsp','ItemsReInsp','DescReInsp',
            'CritReInsp','DescAgrupReInsp','CodReInsp','Clasificacion'
        ]
    else:
        # Formato B: sin columnas de referencia numerica
        cols = [
            'Mes','Responsable','Contrato','Linea','Vehiculo','Modulo','MR',
            'Modelo','Servicio','Fecha','NroInforme','SistemaUnidad',
            'SistemaAmpliado','Item1','Item2','Descripcion',
            'Criticidad','DescAgrupada','CritAmpliado','CodItem',
            'FechaReInsp','NroReInsp','SistUnitReInsp','SistAmpReInsp',
            'ItemsReInsp','DescReInsp','CritReInsp','DescAgrupReInsp','CodReInsp',
            'Clasificacion'
        ]
        # Agregar columnas de referencia vacias para que el resto del codigo funcione
        # (df_out quedara vacio, tab de desvios mostrara mensaje informativo)

    def filas_a_df(rows):
        if not rows:
            return pd.DataFrame(columns=cols)
        n, n_cols = len(rows[0]), len(cols)
        if n == n_cols:
            return pd.DataFrame(rows, columns=cols)
        elif n < n_cols:
            rows_pad = [list(r) + [None] * (n_cols - n) for r in rows]
            return pd.DataFrame(rows_pad, columns=cols)
        else:
            return pd.DataFrame([r[:n_cols] for r in rows], columns=cols)

    dfs = [filas_a_df(rows1)]
    if rows2:
        dfs.append(filas_a_df(rows2))
    df = pd.concat(dfs, ignore_index=True)

    # Normalizar columna Mes (unificar mayusculas/minusculas)
    if 'Mes' in df.columns:
        df['Mes'] = df['Mes'].astype(str).str.strip().str.upper()

    # Limpiar columna Clasificacion: eliminar fórmulas Excel no calculadas
    if 'Clasificacion' in df.columns:
        df['Clasificacion'] = df['Clasificacion'].apply(
            lambda x: None if (pd.isna(x) or str(x).strip().startswith('=')) else str(x).strip()
        )

    # ── 7. Calcular desvios (solo si hay columnas de referencia) ──
    if tiene_refs and all(c in df.columns for c in ['RefMin','RefMax','Relevado1','Relevado2']):
        df_ref = df[df['RefMin'].notna() | df['RefMax'].notna()].copy()
        # R1_num y R2_num mantienen el primer valor (para compatibilidad)
        df_ref['R1_num']     = df_ref['Relevado1'].apply(parse_valor)
        df_ref['R2_num']     = df_ref['Relevado2'].apply(parse_valor)
        df_ref['RefMin_num'] = pd.to_numeric(df_ref['RefMin'], errors='coerce')
        df_ref['RefMax_num'] = pd.to_numeric(df_ref['RefMax'], errors='coerce')
        # calcular_desvio evalúa TODOS los sub-valores separados por / o _
        df_ref[['ValorRelevado','Desviacion']] = df_ref.apply(
            lambda r: pd.Series(calcular_desvio(r)), axis=1
        )
        df_out = df_ref[df_ref['Desviacion'].notna() & (df_ref['Desviacion'] != 0)].copy()
    else:
        # Formato sin referencias: df_out vacio con columnas esperadas
        df_out = pd.DataFrame(columns=list(df.columns) + ['ValorRelevado','Desviacion'])

    return df, df_out


def kpi(label, value, variant="default"):
    return f"""
    <div class="kpi-card {variant}">
        <div class="kpi-value">{value}</div>
        <div class="kpi-label">{label}</div>
    </div>"""




# ─────────────────────────────────────────────
# GENERACIÓN DE WORD (descarga)
# ─────────────────────────────────────────────

def generar_word(df, df_out, config=None):
    from docx.shared import Cm, Mm, Twips
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH as WALIGN
    from datetime import date

    if config is None:
        config = {}

    inc_crit    = config.get("inc_crit",    True)
    inc_top15   = config.get("inc_top15",   True)
    inc_desvios = config.get("inc_desvios", True)
    inc_detalle = config.get("inc_detalle", True)
    inc_concl   = config.get("inc_concl",   True)
    graficos    = config.get("graficos",    {})   # dict key→PNG bytes
    hdr_codigo  = config.get("codigo",  "") or ""
    hdr_version = config.get("version", "v1.0") or "v1.0"
    hdr_linea   = config.get("linea",   "") or ""
    hdr_subger  = config.get("subger",  "Sub Gerencia de Programación y Seguimiento de Mantenimiento de Material Rodante (SPySM)") or ""
    logo_bytes  = config.get("logo",    None)

    doc = Document()

    # ── Página A4, márgenes 15mm ──
    section = doc.sections[0]
    section.page_width      = Mm(210)
    section.page_height     = Mm(297)
    section.left_margin     = Mm(15)
    section.right_margin    = Mm(15)
    section.top_margin      = Mm(35)
    section.bottom_margin   = Mm(15)
    section.header_distance = Mm(5)

    # Ancho de contenido: 210 - 30 = 180mm → DXA
    PAGE_W = int(180 * 56.69)   # 10204 DXA

    # ── Estilos base ──
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    # ── XML helpers ──
    def set_shd(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:shd")):
            tcPr.remove(old)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

    def set_cell_w(cell, w_dxa):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        for old in tcPr.findall(qn("w:tcW")):
            tcPr.remove(old)
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"),    str(w_dxa))
        tcW.set(qn("w:type"), "dxa")
        tcPr.insert(0, tcW)

    def set_tbl_w(tbl, w_dxa):
        tblPr = tbl._tbl.tblPr
        for old in tblPr.findall(qn("w:tblW")):
            tblPr.remove(old)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"),    str(w_dxa))
        tblW.set(qn("w:type"), "dxa")
        tblPr.append(tblW)

    def remove_borders(tbl):
        tblPr = tbl._tbl.tblPr
        for old in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(old)
        b = OxmlElement("w:tblBorders")
        for side in ["top","left","bottom","right","insideH","insideV"]:
            s = OxmlElement(f"w:{side}")
            s.set(qn("w:val"),    "none")
            s.set(qn("w:sz"),     "0")
            s.set(qn("w:space"),  "0")
            s.set(qn("w:color"),  "auto")
            b.append(s)
        tblPr.append(b)

    def cell_text(cell, text, bold=False, size=9, color="000000",
                  align=WD_ALIGN_PARAGRAPH.CENTER, italic=False):
        p = cell.paragraphs[0]
        p.clear()
        p.alignment = align
        r = p.add_run(str(text) if pd.notna(text) else "")
        r.bold   = bold
        r.italic = italic
        r.font.size = Pt(size)
        r.font.name = "Arial"
        r.font.color.rgb = RGBColor.from_string(color)

    # ────────────────────────────────────────
    # ENCABEZADO
    # ────────────────────────────────────────
    header = section.header

    # Borrar párrafos por defecto
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    # Tabla principal: 1 fila, cols dinámicas
    # Si hay logo: [logo | datos]  sino: [texto empresa | datos]
    W_LEFT  = int(PAGE_W * 0.65)
    W_RIGHT = PAGE_W - W_LEFT

    htbl = header.add_table(rows=1, cols=2, width=Mm(180))
    htbl.style = "Table Grid"
    remove_borders(htbl)
    set_tbl_w(htbl, PAGE_W)

    # Altura de fila ~2.5cm
    tr = htbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),    str(int(2.5 * 567)))
    trH.set(qn("w:hRule"),  "atLeast")
    trPr.append(trH)

    c_left  = htbl.cell(0, 0)
    c_right = htbl.cell(0, 1)
    set_cell_w(c_left,  W_LEFT)
    set_cell_w(c_right, W_RIGHT)

    # Celda izquierda — logo o texto
    if logo_bytes:
        set_shd(c_left, "FFFFFF")
        c_left.paragraphs[0].clear()
        c_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_img = c_left.paragraphs[0].add_run()
        run_img.add_picture(io.BytesIO(logo_bytes), height=Cm(2.2))
    else:
        set_shd(c_left, "1F3864")
        cell_text(c_left, "TRENES ARGENTINOS — PISE",
                  bold=True, size=13, color="FFFFFF",
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # Celda derecha — sub-tabla con campos
    set_shd(c_right, "EBF3FB")
    # Vaciar párrafos por defecto de la celda
    for p in list(c_right.paragraphs):
        p._element.getparent().remove(p._element)

    fields = [
        ("Código:",      hdr_codigo  or "___________"),
        ("Versión:",     hdr_version),
        ("Fecha:",       date.today().strftime("%d/%m/%Y")),
        ("Línea:",       hdr_linea   or "___________"),
        ("Subgerencia:", hdr_subger),
    ]

    # Una sola sub-tabla con 5 filas
    sub   = c_right.add_table(rows=len(fields), cols=2)
    remove_borders(sub)
    set_tbl_w(sub, W_RIGHT)
    sub_w = W_RIGHT // 2

    def set_sub_row_height(row, height_cm):
        tr   = row._tr
        trPr = tr.get_or_add_trPr()
        trH  = OxmlElement("w:trHeight")
        trH.set(qn("w:val"),    str(int(height_cm * 567)))
        trH.set(qn("w:hRule"),  "exact")
        trPr.append(trH)

    for i, (lbl, val) in enumerate(fields):
        last = (i == len(fields) - 1)
        if last:
            # Última fila: subgerencia fusionada en ambas columnas
            merged = sub.cell(i, 0).merge(sub.cell(i, 1))
            set_cell_w(merged, W_RIGHT)
            set_shd(merged, "D5E8F0")
            cell_text(merged, f"{lbl} {val}",
                      bold=False, size=7, color="1F3864",
                      align=WD_ALIGN_PARAGRAPH.LEFT, italic=True)
            set_sub_row_height(sub.rows[i], 0.55)  # subgerencia un poco más alta
        else:
            lc = sub.cell(i, 0)
            vc = sub.cell(i, 1)
            set_cell_w(lc, sub_w)
            set_cell_w(vc, sub_w)
            set_shd(lc, "EBF3FB")
            set_shd(vc, "EBF3FB")
            cell_text(lc, lbl, bold=True,  size=8, color="1F3864", align=WD_ALIGN_PARAGRAPH.LEFT)
            cell_text(vc, val, bold=False, size=8, color="333333", align=WD_ALIGN_PARAGRAPH.LEFT)
            set_sub_row_height(sub.rows[i], 0.42)  # filas de datos compactas

    # ── PIE DE PÁGINA ──
    footer = section.footer
    for p in list(footer.paragraphs):
        p.clear()
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_f = fp.add_run(f"SPySM  —  {date.today().strftime('%d/%m/%Y')}  —  Pág. ")
    r_f.font.size = Pt(8)
    r_f.font.name = "Arial"
    r_f.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # ────────────────────────────────────────
    # TABLE HELPER — ancho exacto PAGE_W
    # ────────────────────────────────────────
    def table_from_df(dataframe):
        n   = len(dataframe.columns)
        # Distribuir ancho: descripcion toma más espacio si existe
        col_names = list(dataframe.columns)
        wide_cols = {"Descripción", "Descripcion", "Categoría", "Categoria",
                     "Descripción Agrupada", "Falla"}
        base_w = PAGE_W // n
        widths = []
        wide_total = sum(1 for c in col_names if c in wide_cols)
        narrow_w   = max(800, int(PAGE_W * 0.07))
        if wide_total:
            wide_w = (PAGE_W - narrow_w * (n - wide_total)) // wide_total
        else:
            wide_w = base_w
        for c in col_names:
            widths.append(wide_w if c in wide_cols else narrow_w if n > 3 else base_w)
        # Normalizar para que sumen exactamente PAGE_W
        total = sum(widths)
        if total != PAGE_W:
            widths[-1] += PAGE_W - total

        t = doc.add_table(rows=1, cols=n)
        t.style   = "Table Grid"
        t.autofit = False
        set_tbl_w(t, PAGE_W)

        # Encabezado
        for i, (cname, w) in enumerate(zip(col_names, widths)):
            cell = t.rows[0].cells[i]
            set_cell_w(cell, w)
            set_shd(cell, "1F3864")
            cell_text(cell, cname, bold=True, size=9,
                      color="FFFFFF", align=WD_ALIGN_PARAGRAPH.CENTER)

        # Datos con filas alternas
        for ri, (_, row) in enumerate(dataframe.iterrows()):
            fill = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
            tr   = t.add_row()
            for i, (val, w) in enumerate(zip(row, widths)):
                cell = tr.cells[i]
                set_cell_w(cell, w)
                set_shd(cell, fill)
                txt  = "" if pd.isna(val) else str(val)
                cell_text(cell, txt, size=9,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
        return t

    # ── Numerador automático ──
    sec_num = [0]
    def next_sec(title):
        sec_num[0] += 1
        h = doc.add_heading(f"{sec_num[0]}. {title}", level=1)
        h.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        h.runs[0].font.name = "Arial"

    # ────────────────────────────────────────
    # TÍTULO
    # ────────────────────────────────────────
    doc.add_paragraph()
    tit = doc.add_heading("Informe de Fallas y Desvíos", 0)
    tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tit.runs[0].font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    tit.runs[0].font.name = "Arial"

    p_meta = doc.add_paragraph(
        f"Período: {df['Mes'].dropna().nunique()} meses  "
        f"|  Vehículos: {df['Vehiculo'].dropna().nunique()}  "
        f"|  Total observaciones: {len(df)}"
    )
    p_meta.runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # ────────────────────────────────────────
    # SECCIONES
    # ────────────────────────────────────────
    if inc_crit:
        next_sec("Distribución por Criticidad")
        d = df["CritAmpliado"].value_counts().reset_index()
        d.columns = ["Criticidad", "Cantidad"]
        d["Porcentaje"] = (d["Cantidad"] / len(df) * 100).round(1).astype(str) + "%"
        table_from_df(d)
        doc.add_paragraph()

    if inc_top15:
        next_sec("Top 15 Tipos de Falla")
        d = df["DescAgrupada"].value_counts().head(15).reset_index()
        d.columns = ["Descripción", "Cantidad"]
        d["% total"] = (d["Cantidad"] / len(df) * 100).round(1).astype(str) + "%"
        table_from_df(d)
        doc.add_paragraph()

    tiene_desv = not df_out.empty and "Desviacion" in df_out.columns

    if tiene_desv and inc_desvios:
        next_sec("Resumen de Desvíos por Categoría")
        d = df_out.groupby("DescAgrupada").agg(
            Cantidad     =("Desviacion", "count"),
            Desv_Promedio=("Desviacion", lambda x: round(x.mean(),   2)),
            Desv_Max_Abs =("Desviacion", lambda x: round(x.abs().max(), 2)),
            Criticos     =("Criticidad", lambda x: (x == "C").sum())
        ).reset_index().sort_values("Cantidad", ascending=False)
        d.columns = ["Categoría", "Cantidad", "Desvío Prom.", "Desvío Máx.", "Críticos"]
        table_from_df(d)
        doc.add_paragraph()

    if tiene_desv and inc_detalle:
        next_sec("Detalle — Observaciones Fuera de Parámetro")
        cols_d  = ["Vehiculo","Mes","SistemaUnidad","Descripcion","Criticidad"]
        rnames  = ["Vehículo","Mes","Sistema","Descripción","Crit."]
        for col in ["RefMin","RefMax","ValorRelevado","Desviacion"]:
            if col in df_out.columns:
                cols_d.insert(-1, col)
                rnames.insert(-1, {"RefMin":"Ref Min","RefMax":"Ref Max",
                                   "ValorRelevado":"Relevado","Desviacion":"Desvío"}[col])
        det = df_out[cols_d].copy()
        det.columns = rnames
        if "Desvío"   in det.columns: det["Desvío"]   = det["Desvío"].round(2)
        if "Relevado" in det.columns: det["Relevado"] = det["Relevado"].round(1)
        table_from_df(det)
        doc.add_paragraph()

    if not tiene_desv and (inc_desvios or inc_detalle):
        next_sec("Observaciones sin valores de referencia")
        p = doc.add_paragraph(
            "Este archivo no contiene columnas de valores de referencia paramétrica. "
            "El análisis de desvíos no está disponible para este formato."
        )
        p.runs[0].font.size = Pt(10)

    # ────────────────────────────────────────
    # CONCLUSIONES DINÁMICAS
    # ────────────────────────────────────────
    if inc_concl:
        next_sec("Conclusiones y Observaciones Generales")
        doc.add_paragraph(
            "Del análisis de la totalidad de las inspecciones estáticas realizadas "
            "durante el período se extraen las siguientes conclusiones:"
        )
        doc.add_paragraph()
        conclusiones = []
        n_tot = len(df)

        sv  = df["SistemaUnidad"].value_counts()
        conclusiones.append(
            f"El sistema de {SISTEMA_LABELS.get(sv.index[0], sv.index[0])} ({sv.index[0]}) "
            f"concentra la mayor cantidad de observaciones con {int(sv.iloc[0])} casos "
            f"({round(sv.iloc[0]/n_tot*100,1)}% del total)."
        )

        df["_unidad"] = df["Modulo"].apply(
            lambda x: None if (x is None or str(x).strip() in ("","0","nan")) else str(x).strip()
        ).fillna(df["Vehiculo"].astype(str).str.strip())
        uv = df["_unidad"].value_counts()
        conclusiones.append(
            f"La unidad con mayor cantidad de observaciones es {uv.index[0]} "
            f"con {int(uv.iloc[0])} casos, siendo candidata prioritaria para revisión integral."
        )

        if "Clasificacion" in df.columns and df["Clasificacion"].notna().any():
            cv = df["Clasificacion"].value_counts()
            conclusiones.append(
                f"La clasificación de falla predominante es '{cv.index[0]}' con {int(cv.iloc[0])} casos "
                f"({round(cv.iloc[0]/n_tot*100,1)}%), seguida por "
                f"'{cv.index[1] if len(cv)>1 else '-'}' con {int(cv.iloc[1]) if len(cv)>1 else 0} casos."
            )

        tc = int((df["Criticidad"]=="C").sum())
        tr = int((df["Criticidad"]=="R").sum())
        conclusiones.append(
            f"Las observaciones de criticidad alta (Crítico + Rechazado) representan el "
            f"{round((tc+tr)/n_tot*100,1)}% del total ({tc} críticas y {tr} rechazadas)."
        )

        fv = df["DescAgrupada"].value_counts()
        conclusiones.append(
            f"La falla más recurrente es '{fv.index[0]}' con {int(fv.iloc[0])} casos "
            f"({round(fv.iloc[0]/n_tot*100,1)}% del total)."
        )

        if tiene_desv and not df_out.empty:
            dg = df_out.groupby("DescAgrupada").agg(
                Cant=("Desviacion","count"),
                Prom=("Desviacion", lambda x: round(x.abs().mean(), 2)),
                Max =("Desviacion", lambda x: round(x.abs().max(),  2)),
            ).sort_values("Max", ascending=False)
            fm   = df_out.loc[df_out["Desviacion"].abs().idxmax()]
            dval = round(float(fm["Desviacion"]), 2)
            conclusiones.append(
                f"La categoría con mayor desvío paramétrico es '{dg.index[0]}' "
                f"({int(dg.iloc[0]['Cant'])} casos, desvío prom. {dg.iloc[0]['Prom']}, "
                f"máx. {dg.iloc[0]['Max']} unidades). "
                f"Caso extremo: {fm['Vehiculo']} — valor relevado "
                f"{round(float(fm['ValorRelevado']),1) if pd.notna(fm['ValorRelevado']) else '-'}, "
                f"{'por encima del máximo' if dval>0 else 'por debajo del mínimo'} "
                f"en {abs(dval)} unidades."
            )

        for idx, texto in enumerate(conclusiones, 1):
            p   = doc.add_paragraph(style="List Number")
            run = p.add_run(texto)
            run.font.size = Pt(10)
            run.font.name = "Arial"

    # ── GRÁFICOS ──
    def insert_graph(key, caption):
        png = graficos.get(key)
        if not png:
            return
        next_sec(f"Gráfico — {caption}")
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_img = p_img.add_run()
        run_img.add_picture(io.BytesIO(png), width=Mm(170))
        doc.add_paragraph()

    insert_graph('torta',      'Distribución por Criticidad')
    insert_graph('sistemas',   'Observaciones por Sistema')
    insert_graph('mensual',    'Evolución Mensual y Ratio Obs/Vehículo')
    insert_graph('top15',      'Top 15 Tipos de Falla')
    insert_graph('desvios',    'Desvíos: Distribución por Categoría')
    insert_graph('pareto_fr',  'Pareto — Fuera de Rango')
    insert_graph('pareto_aus', 'Pareto — Ausencia de Elementos')
    insert_graph('pareto_mal', 'Pareto — Mal Estado')
    insert_graph('explorador', 'Explorador Libre')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ─────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────

with st.sidebar:
    st.markdown("## 🔧 Análisis de Mantenimiento")
    st.markdown("---")
    uploaded = st.file_uploader(
        "Subir archivo Excel",
        type=["xlsx"],
        help="Archivo de seguimiento de inspecciones estáticas"
    )
    st.markdown("---")
    st.markdown("""
    **Cómo usar:**
    1. Subí el archivo `.xlsx`
    2. Explorá los paneles
    3. Filtrá por mes, sistema o vehículo
    4. Descargá el informe Word
    """)
    st.markdown("---")
    st.caption("Análisis automático de desvíos · Material Rodante")


# ─────────────────────────────────────────────
# PANTALLA PRINCIPAL
# ─────────────────────────────────────────────

st.title("🔧 Análisis de Informes de Mantenimiento")

if uploaded is None:
    st.markdown("""
    <div style="
        background: linear-gradient(135deg,#1a2235,#1e2a3a);
        border: 1px dashed #2a3a50;
        border-radius: 12px;
        padding: 60px;
        text-align: center;
        margin-top: 40px;
    ">
        <div style="font-size:3rem;margin-bottom:16px">📂</div>
        <div style="color:#4fc3f7;font-size:1.3rem;font-weight:700;margin-bottom:8px">
            Subí tu archivo Excel para comenzar
        </div>
        <div style="color:#546e7a;font-size:0.9rem">
            Usá el panel izquierdo para cargar el archivo de inspecciones
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.stop()

# ── Cargar y procesar ──
with st.spinner("Procesando archivo..."):
    df, df_out = cargar_y_analizar(uploaded.read())

total_obs        = len(df)
total_veh        = df['Vehiculo'].dropna().nunique()
total_normal     = int((df['Criticidad'] == 'N').sum())
total_crit       = int((df['Criticidad'] == 'C').sum())
total_rech       = int((df['Criticidad'] == 'R').sum())
total_corregidas = int((df['Criticidad'] == 'O').sum())
total_nrc        = int((df['Criticidad'] == 'NRC').sum())
total_desv       = len(df_out)
pct_alta         = round((total_crit + total_rech) / total_obs * 100, 1) if total_obs > 0 else 0

# ─────────────────────────────────────────────
# KPIs — fila 1: volumen general
# ─────────────────────────────────────────────
st.markdown('<div class="section-header">Resumen del período</div>', unsafe_allow_html=True)

c1, c2, c3, c4 = st.columns(4)
c1.markdown(kpi("Total Observaciones", total_obs),          unsafe_allow_html=True)
c2.markdown(kpi("Vehículos Inspeccionados", total_veh),     unsafe_allow_html=True)
c3.markdown(kpi("Sin Observaciones", total_nrc),      unsafe_allow_html=True)
c4.markdown(kpi("% Criticidad Alta (Rechazo + Critico)", f"{pct_alta}%", "danger"), unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# KPIs — fila 2: desglose por criticidad
# ─────────────────────────────────────────────
d1, d2, d3, d4 = st.columns(4)
d1.markdown(kpi("Normales",          total_normal,     "success"), unsafe_allow_html=True)
d2.markdown(kpi("Corregidas en Inspección", total_corregidas, "success"), unsafe_allow_html=True)
d3.markdown(kpi("Críticas",          total_crit,       "danger"),  unsafe_allow_html=True)
d4.markdown(kpi("Rechazadas",        total_rech,       "warning"), unsafe_allow_html=True)

st.markdown("<br>", unsafe_allow_html=True)

# ─────────────────────────────────────────────
# TABS PRINCIPALES
# ─────────────────────────────────────────────
tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs([
    "📊  Resúmenes",
    "📈  Gráficos",
    "🔍  Desvíos Detallados",
    "🧩  Análisis por Clasificación",
    "🔬  Explorador Libre",
    "📉  Indicadores",
    "📄  Exportar Informe"
])

PLOTLY_THEME = dict(
    paper_bgcolor='rgba(0,0,0,0)',
    plot_bgcolor='rgba(26,34,53,0.6)',
    font=dict(family='IBM Plex Sans', color='#b0bec5'),
    colorway=['#4fc3f7','#81d4fa','#ef5350','#ffa726','#66bb6a','#ab47bc','#26c6da'],
)
# Estilos de ejes reutilizables (separados para no generar claves duplicadas en update_layout)
AXIS_STYLE = dict(gridcolor='#1e2a3a', linecolor='#2a3a50')


# ── TAB 1: RESÚMENES ──
with tab1:
    col_a, col_b = st.columns(2)

    with col_a:
        st.markdown("#### Criticidad")
        crit_df = df['CritAmpliado'].value_counts().reset_index()
        crit_df.columns = ['Criticidad','Cantidad']
        crit_df['Porcentaje'] = (crit_df['Cantidad'] / total_obs * 100).round(1).astype(str) + '%'
        st.dataframe(crit_df, use_container_width=True, hide_index=True)

        st.markdown("#### Por Modelo")
        mod_df = df['Modelo'].value_counts().reset_index()
        mod_df.columns = ['Modelo','Observaciones']
        mod_df['%'] = (mod_df['Observaciones'] / total_obs * 100).round(1).astype(str) + '%'
        st.dataframe(mod_df, use_container_width=True, hide_index=True)

    with col_b:
        st.markdown("#### Por Sistema de Unidad")
        sist_df = df['SistemaUnidad'].value_counts().reset_index()
        sist_df.columns = ['Código','Cantidad']
        sist_df['Sistema'] = sist_df['Código'].map(SISTEMA_LABELS).fillna(sist_df['Código'])
        sist_df['%'] = (sist_df['Cantidad'] / total_obs * 100).round(1).astype(str) + '%'
        st.dataframe(sist_df[['Código','Sistema','Cantidad','%']], use_container_width=True, hide_index=True)

    st.markdown("#### Top 15 Tipos de Falla")
    top15 = df['DescAgrupada'].value_counts().head(15).reset_index()
    top15.columns = ['Descripción Agrupada','Cantidad']
    top15['% del total'] = (top15['Cantidad'] / total_obs * 100).round(1).astype(str) + '%'
    st.dataframe(top15, use_container_width=True, hide_index=True)

    # ── TOP 10 POR TIPO DE MR ──
    st.markdown("---")
    st.markdown("#### Top 10 por Tipo de Material Rodante")

    MR_CONFIG = {
        'LOC' : {'label': 'Locomotoras (LOC)',        'servicios': ['LD', 'PERIFERICO', 'AMBA']},
        'CCRR': {'label': 'Coches Remolcados (CCRR)', 'servicios': ['LD', 'PERIFERICO']},
        'CCEE': {'label': 'Coches Eléctricos (CCEE)', 'servicios': []},
        'CCMM': {'label': 'Coche Motor (CCMM)',       'servicios': ['PERIFERICO', 'AMBA']},
    }

    def top10_unidad(df_base):
        """Aplica regla Modulo > Vehiculo y devuelve top 10."""
        df_base = df_base.copy()
        df_base['_unidad'] = df_base['Modulo'].apply(
            lambda x: None if (x is None or str(x).strip() in ('', '0', 'nan')) else str(x).strip()
        ).fillna(df_base['Vehiculo'].astype(str).str.strip())
        result = df_base['_unidad'].value_counts().head(10).reset_index()
        result.columns = ['Unidad', 'Observaciones']
        return result

    # Fila 1: LOC y CCRR
    row_t1, row_t2 = st.columns(2)

    with row_t1:
        cfg = MR_CONFIG['LOC']
        st.markdown(f"##### {cfg['label']}")
        df_mr = df[df['MR'] == 'LOC']
        if df_mr.empty:
            st.caption("NO SE INSPECCIONÓ ESTE TIPO DE MR")
        else:
            servicios_disp = sorted(df_mr['Servicio'].dropna().unique().tolist())
            sel_srv = st.multiselect("Servicio", servicios_disp,
                                     default=servicios_disp, key="srv_loc")
            df_fil = df_mr[df_mr['Servicio'].isin(sel_srv)] if sel_srv else df_mr
            st.dataframe(top10_unidad(df_fil), use_container_width=True, hide_index=True)

    with row_t2:
        cfg = MR_CONFIG['CCRR']
        st.markdown(f"##### {cfg['label']}")
        df_mr = df[df['MR'] == 'CCRR']
        if df_mr.empty:
            st.caption("NO SE INSPECCIONÓ ESTE TIPO DE MR")
        else:
            servicios_disp = sorted(df_mr['Servicio'].dropna().unique().tolist())
            sel_srv = st.multiselect("Servicio", servicios_disp,
                                     default=servicios_disp, key="srv_ccrr")
            df_fil = df_mr[df_mr['Servicio'].isin(sel_srv)] if sel_srv else df_mr
            st.dataframe(top10_unidad(df_fil), use_container_width=True, hide_index=True)

    # Fila 2: CCEE y CCMM
    row_t3, row_t4 = st.columns(2)

    with row_t3:
        cfg = MR_CONFIG['CCEE']
        st.markdown(f"##### {cfg['label']}")
        df_mr = df[df['MR'] == 'CCEE']
        if df_mr.empty:
            st.caption("NO SE INSPECCIONÓ ESTE TIPO DE MR")
        else:
            # CCEE solo tiene AMBA, no necesita filtro
            st.dataframe(top10_unidad(df_mr), use_container_width=True, hide_index=True)

    with row_t4:
        cfg = MR_CONFIG['CCMM']
        st.markdown(f"##### {cfg['label']}")
        df_mr = df[df['MR'] == 'CCMM']
        if df_mr.empty:
            st.caption("NO SE INSPECCIONÓ ESTE TIPO DE MR")
        else:
            servicios_disp = sorted(df_mr['Servicio'].dropna().unique().tolist())
            sel_srv = st.multiselect("Servicio", servicios_disp,
                                     default=servicios_disp, key="srv_ccmm")
            df_fil = df_mr[df_mr['Servicio'].isin(sel_srv)] if sel_srv else df_mr
            st.dataframe(top10_unidad(df_fil), use_container_width=True, hide_index=True)


# ── TAB 2: GRÁFICOS ──
with tab2:
    row1_l, row1_r = st.columns(2)

    # Torta criticidad
    with row1_l:
        st.markdown("#### Distribución por Criticidad")
        crit_counts = df['CritAmpliado'].value_counts()
        fig_pie = go.Figure(go.Pie(
            labels=crit_counts.index,
            values=crit_counts.values,
            hole=0.55,
            marker=dict(colors=['#4fc3f7','#ef5350','#ffa726','#66bb6a','#ab47bc']),
            textfont=dict(size=13)
        ))
        fig_pie.update_layout(**PLOTLY_THEME, margin=dict(t=10,b=10,l=10,r=10), height=300,
                              legend=dict(orientation='h', y=-0.1),
                              xaxis=AXIS_STYLE, yaxis=AXIS_STYLE)
        st.plotly_chart(fig_pie, use_container_width=True, key="chart_pie")

    # Barras sistemas
    with row1_r:
        st.markdown("#### Observaciones por Sistema")
        sist_counts = df['SistemaUnidad'].value_counts().reset_index()
        sist_counts.columns = ['Sistema','Cantidad']
        sist_counts['Label'] = sist_counts['Sistema'].map(SISTEMA_LABELS).fillna(sist_counts['Sistema'])
        fig_sist = px.bar(sist_counts, x='Cantidad', y='Label', orientation='h',
                          color='Cantidad', color_continuous_scale='Blues',
                          text='Cantidad')
        fig_sist.update_traces(textposition='outside', textfont=dict(color='#b0bec5', size=12))
        fig_sist.update_layout(**PLOTLY_THEME, margin=dict(t=10,b=10,r=80,l=10), height=300,
                               xaxis=dict(range=[0, sist_counts['Cantidad'].max() * 1.15], **AXIS_STYLE),
                               yaxis=dict(autorange='reversed', **AXIS_STYLE),
                               coloraxis_showscale=False)
        st.plotly_chart(fig_sist, use_container_width=True, key="chart_sist")

    # Evolución mensual + ratio obs/vehículo
    st.markdown("#### Evolución Mensual de Observaciones")
    monthly = []
    for m in MONTH_ORDER:
        df_m = df[df['Mes'].str.upper() == m]
        cnt  = len(df_m)
        vehs = df_m['Vehiculo'].dropna().nunique()
        if cnt > 0:
            monthly.append({
                'Mes': m,
                'Observaciones': cnt,
                'Vehículos': vehs,
                'Ratio Obs/Veh': round(cnt / vehs, 2) if vehs > 0 else 0
            })
    monthly_df = pd.DataFrame(monthly)
    if not monthly_df.empty:
        fig_line = go.Figure()

        # Barras de observaciones
        fig_line.add_trace(go.Bar(
            x=monthly_df['Mes'], y=monthly_df['Observaciones'],
            name='Observaciones',
            marker_color='rgba(79,195,247,0.3)',
            text=monthly_df['Observaciones'],
            textposition='outside',
            textfont=dict(color='#4fc3f7', size=11),
            yaxis='y'
        ))

        # Línea de vehículos inspeccionados
        fig_line.add_trace(go.Scatter(
            x=monthly_df['Mes'], y=monthly_df['Vehículos'],
            name='Vehículos inspeccionados',
            mode='lines+markers',
            line=dict(color='#66bb6a', width=2, dash='dot'),
            marker=dict(size=7, color='#66bb6a'),
            yaxis='y'
        ))

        # Línea ratio obs/vehículo — eje secundario
        fig_line.add_trace(go.Scatter(
            x=monthly_df['Mes'], y=monthly_df['Ratio Obs/Veh'],
            name='Ratio Obs/Veh',
            mode='lines+markers+text',
            line=dict(color='#ffa726', width=2),
            marker=dict(size=8, color='#ffa726'),
            text=monthly_df['Ratio Obs/Veh'].astype(str),
            textposition='top center',
            textfont=dict(color='#ffa726', size=10),
            yaxis='y2'
        ))

        fig_line.update_layout(
            **PLOTLY_THEME,
            height=360,
            margin=dict(t=20, b=20, l=10, r=60),
            xaxis=AXIS_STYLE,
            yaxis=dict(title='Cantidad', **AXIS_STYLE),
            yaxis2=dict(title='Ratio Obs/Veh', overlaying='y', side='right',
                        gridcolor='#1e2a3a', linecolor='#2a3a50'),
            legend=dict(orientation='h', y=1.08),
            barmode='overlay'
        )
        st.plotly_chart(fig_line, use_container_width=True, key="chart_line")

    # Top fallas barras
    st.markdown("#### Top 15 Tipos de Falla")
    top15_plot = df['DescAgrupada'].value_counts().head(15).reset_index()
    top15_plot.columns = ['Falla','Cantidad']
    fig_top = px.bar(top15_plot, x='Cantidad', y='Falla', orientation='h',
                     color='Cantidad', color_continuous_scale='Blues_r',
                     text='Cantidad')
    fig_top.update_traces(textposition='outside', textfont=dict(color='#b0bec5', size=12))
    fig_top.update_layout(**PLOTLY_THEME, height=420, margin=dict(t=10,b=10,r=80,l=10),
                          xaxis=dict(range=[0, top15_plot['Cantidad'].max() * 1.15], **AXIS_STYLE),
                          yaxis=dict(autorange='reversed', **AXIS_STYLE),
                          coloraxis_showscale=False)
    st.plotly_chart(fig_top, use_container_width=True, key="chart_top")

    # Desvíos por categoría
    if not df_out.empty and 'Desviacion' in df_out.columns:
        st.markdown("#### Desvíos: Distribución por Categoría")
        st.caption("Barras = cantidad de casos · Líneas = desvío máx/prom/mín sobre eje derecho (mismas unidades que el parámetro)")
        desv_plot = df_out.groupby('DescAgrupada').agg(
            Cantidad  =('Desviacion', 'count'),
            Desv_Max  =('Desviacion', lambda x: round(x.abs().max(),  2)),
            Desv_Prom =('Desviacion', lambda x: round(x.abs().mean(), 2)),
            Desv_Min  =('Desviacion', lambda x: round(x.abs().min(),  2)),
        ).reset_index().sort_values('Cantidad', ascending=False)

        # Etiquetas combinadas para tooltips
        desv_plot['Etiqueta'] = desv_plot.apply(
            lambda r: f"Máx: {r['Desv_Max']}  |  Prom: {r['Desv_Prom']}  |  Mín: {r['Desv_Min']}", axis=1
        )

        fig_desv = go.Figure()

        # Barras cantidad
        fig_desv.add_trace(go.Bar(
            x=desv_plot['DescAgrupada'], y=desv_plot['Cantidad'],
            name='Cantidad de casos',
            marker_color='#4fc3f7',
            text=desv_plot['Cantidad'],
            textposition='outside',
            textfont=dict(color='#4fc3f7', size=11),
            yaxis='y'
        ))

        # Línea desvío máximo
        fig_desv.add_trace(go.Scatter(
            x=desv_plot['DescAgrupada'], y=desv_plot['Desv_Max'],
            name='Desvío Máx.',
            mode='lines+markers+text',
            line=dict(color='#ef5350', width=2),
            marker=dict(size=8, color='#ef5350'),
            text=desv_plot['Desv_Max'].astype(str),
            textposition='top center',
            textfont=dict(color='#ef5350', size=10),
            yaxis='y2'
        ))

        # Línea desvío promedio
        fig_desv.add_trace(go.Scatter(
            x=desv_plot['DescAgrupada'], y=desv_plot['Desv_Prom'],
            name='Desvío Prom.',
            mode='lines+markers+text',
            line=dict(color='#ffa726', width=2, dash='dot'),
            marker=dict(size=7, color='#ffa726'),
            text=desv_plot['Desv_Prom'].astype(str),
            textposition='bottom center',
            textfont=dict(color='#ffa726', size=10),
            yaxis='y2'
        ))

        # Línea desvío mínimo
        fig_desv.add_trace(go.Scatter(
            x=desv_plot['DescAgrupada'], y=desv_plot['Desv_Min'],
            name='Desvío Mín.',
            mode='lines+markers',
            line=dict(color='#66bb6a', width=1, dash='dash'),
            marker=dict(size=6, color='#66bb6a'),
            yaxis='y2'
        ))

        fig_desv.update_layout(
            **PLOTLY_THEME, height=440, barmode='group',
            margin=dict(t=20, b=100, l=10, r=70),
            xaxis=dict(tickangle=-35, **AXIS_STYLE),
            yaxis=dict(title='Cantidad de casos', **AXIS_STYLE),
            yaxis2=dict(title='Desvío (unidades del parámetro)', overlaying='y', side='right',
                        gridcolor='#1e2a3a', linecolor='#2a3a50'),
            legend=dict(orientation='h', y=1.06)
        )
        st.plotly_chart(fig_desv, use_container_width=True, key="chart_desv")


# ── TAB 3: DESVÍOS DETALLADOS ──
with tab3:
    if df_out.empty or 'Desviacion' not in df_out.columns:
        st.info(
            "Este archivo no contiene columnas de valores de referencia (RefMin / RefMax / Relevado). "
            "El análisis de desvíos paramétricos no está disponible para este formato. "
            "Las observaciones cualitativas se muestran en el tab **Resúmenes**."
        )
    else:
        st.markdown("#### Filtros")
        fc1, fc2, fc3, fc4, fc5 = st.columns(5)

        meses_disponibles    = sorted(df_out['Mes'].dropna().unique().tolist())
        sistemas_disponibles = sorted(df_out['SistemaUnidad'].dropna().unique().tolist())
        criticas_opciones    = sorted(df_out['Criticidad'].dropna().unique().tolist())

        # Unidades (Modulo > Vehiculo)
        df_out['_unidad'] = df_out['Modulo'].apply(
            lambda x: None if (x is None or str(x).strip() in ('', '0', 'nan')) else str(x).strip()
        ).fillna(df_out['Vehiculo'].astype(str).str.strip())
        unidades_disponibles = sorted(df_out['_unidad'].dropna().unique().tolist())

        # Descripciones (equivale a rango ref / tipo de medición)
        desc_disponibles = sorted(df_out['DescAgrupada'].dropna().unique().tolist())

        with fc1:
            sel_mes = st.multiselect("Mes", meses_disponibles, default=meses_disponibles,
                                     placeholder="Todos los meses", key="f3_mes")
        with fc2:
            sel_sist = st.multiselect("Sistema", sistemas_disponibles, default=sistemas_disponibles,
                                      placeholder="Todos los sistemas", key="f3_sist")
        with fc3:
            sel_crit = st.multiselect("Criticidad", criticas_opciones, default=criticas_opciones,
                                      placeholder="Todas", key="f3_crit")
        with fc4:
            sel_unidad = st.multiselect("Vehículo / Módulo", unidades_disponibles,
                                        default=unidades_disponibles, placeholder="Todas las unidades",
                                        key="f3_unidad")
        with fc5:
            sel_desc = st.multiselect("Tipo de medición", desc_disponibles,
                                      default=desc_disponibles, placeholder="Todas",
                                      key="f3_desc")

        df_filtrado = df_out[
            df_out['Mes'].isin(sel_mes) &
            df_out['SistemaUnidad'].isin(sel_sist) &
            df_out['Criticidad'].isin(sel_crit) &
            df_out['_unidad'].isin(sel_unidad) &
            df_out['DescAgrupada'].isin(sel_desc)
        ].copy()

        def ref_str(row):
            mn = row['RefMin_num'] if pd.notna(row.get('RefMin_num')) else None
            mx = row['RefMax_num'] if pd.notna(row.get('RefMax_num')) else None
            if mn is not None and mx is not None: return f"{mn} - {mx}"
            if mn is not None: return f">= {mn}"
            if mx is not None: return f"<= {mx}"
            return "-"

        df_filtrado['RefMin_num'] = pd.to_numeric(df_filtrado.get('RefMin'), errors='coerce')
        df_filtrado['RefMax_num'] = pd.to_numeric(df_filtrado.get('RefMax'), errors='coerce')
        df_filtrado['Rango Ref.'] = df_filtrado.apply(ref_str, axis=1)
        df_filtrado['Desvio']     = df_filtrado['Desviacion'].round(2)
        df_filtrado['Relevado']   = df_filtrado['ValorRelevado'].round(1)

        st.markdown(f"**{len(df_filtrado)}** observaciones fuera de parámetro")

        st.markdown("##### Resumen por Categoría")
        resumen = df_filtrado.groupby('DescAgrupada').agg(
            Cantidad=('Desvio','count'),
            Desv_Promedio=('Desvio', lambda x: round(x.mean(),2)),
            Desv_Max=('Desvio', lambda x: round(x.abs().max(),2)),
            Criticos=('Criticidad', lambda x: (x=='C').sum())
        ).reset_index().sort_values('Cantidad', ascending=False)
        resumen.columns = ['Categoria','Cantidad','Desvio Prom.','Desvio Max. (abs)','Criticos']
        st.dataframe(resumen, use_container_width=True, hide_index=True)

        st.markdown("##### Detalle completo")
        cols_show  = ['Vehiculo','Mes','SistemaUnidad','Descripcion','Rango Ref.','Relevado','Desvio','Criticidad']
        rename_map = {'Vehiculo':'Vehiculo','SistemaUnidad':'Sistema','Descripcion':'Descripcion'}
        tabla_det  = df_filtrado[cols_show].rename(columns=rename_map).reset_index(drop=True)

        st.dataframe(
            tabla_det,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Desvio":    st.column_config.NumberColumn(format="%.2f"),
                "Relevado":  st.column_config.NumberColumn(format="%.1f"),
                "Criticidad": st.column_config.TextColumn(),
            }
        )




# ── TAB 4: ANÁLISIS POR CLASIFICACIÓN ──
with tab4:
    tiene_clasif = 'Clasificacion' in df.columns and df['Clasificacion'].notna().any()

    if not tiene_clasif:
        st.info("Este archivo no tiene columna 'Clasificacion'. Agregala en el Excel para habilitar este análisis.")
    else:
        # ── Tabla cruzada Sistema × Criticidad ──
        st.markdown("#### Tabla Cruzada: Sistema × Criticidad")
        st.caption("Cantidad de observaciones por sistema y nivel de criticidad")

        pivot = df.pivot_table(
            index='CritAmpliado',
            columns='SistemaUnidad',
            aggfunc='size',
            fill_value=0
        )
        # Ordenar filas por criticidad
        orden_crit = ['Rechazado','Critico','Normal','Corregida']
        pivot = pivot.reindex([r for r in orden_crit if r in pivot.index])
        st.dataframe(pivot, use_container_width=True)

        st.markdown("---")

        # ── 3 Paretos por Clasificacion ──
        st.markdown("#### Pareto por Clasificación de Falla")
        st.caption("Cada gráfico muestra qué sistemas concentran el 80% de las fallas en cada categoría")

        CLASIF_COLORS = {
            'Fuera de rango':       '#ef5350',
            'Ausencia de elementos':'#ffa726',
            'Mal estado':           '#4fc3f7',
        }

        categorias = ['Fuera de rango', 'Ausencia de elementos', 'Mal estado']
        cols_pareto = st.columns(3)

        for i, cat in enumerate(categorias):
            df_cat = df[df['Clasificacion'].str.strip().str.lower() == cat.lower()].copy()
            with cols_pareto[i]:
                st.markdown(f"##### {cat}")
                if df_cat.empty:
                    st.caption("Sin datos")
                    continue

                # Contar por sistema y calcular % acumulado
                conteo = df_cat['SistemaUnidad'].value_counts().reset_index()
                conteo.columns = ['Sistema', 'Cantidad']
                conteo['Label'] = conteo['Sistema'].map(SISTEMA_LABELS).fillna(conteo['Sistema'])
                conteo['%_acum'] = (conteo['Cantidad'].cumsum() / conteo['Cantidad'].sum() * 100).round(1)

                color = CLASIF_COLORS.get(cat, '#4fc3f7')

                fig_p = go.Figure()

                # Barras
                fig_p.add_trace(go.Bar(
                    x=conteo['Label'],
                    y=conteo['Cantidad'],
                    name='Cantidad',
                    marker_color=color,
                    text=conteo['Cantidad'],
                    textposition='outside',
                    textfont=dict(color='#e8eaf0', size=11),
                    yaxis='y'
                ))

                # Línea acumulada
                fig_p.add_trace(go.Scatter(
                    x=conteo['Label'],
                    y=conteo['%_acum'],
                    name='% Acum.',
                    mode='lines+markers+text',
                    line=dict(color='#ffffff', width=2),
                    marker=dict(size=6, color='#ffffff'),
                    text=[f"{v}%" for v in conteo['%_acum']],
                    textposition='top center',
                    textfont=dict(color='#ffffff', size=10),
                    yaxis='y2'
                ))

                # Línea de referencia 80%
                fig_p.add_hline(
                    y=80, line_dash='dash',
                    line_color='#ffa726', line_width=1,
                    annotation_text='80%',
                    annotation_font_color='#ffa726',
                    yref='y2'
                )

                fig_p.update_layout(
                    **PLOTLY_THEME,
                    height=350,
                    margin=dict(t=30, b=60, l=10, r=50),
                    showlegend=False,
                    xaxis=dict(tickangle=-35, **AXIS_STYLE),
                    yaxis=dict(title='Cantidad', **AXIS_STYLE),
                    yaxis2=dict(
                        title='% Acumulado',
                        overlaying='y', side='right',
                        range=[0, 110],
                        gridcolor='#1e2a3a', linecolor='#2a3a50'
                    ),
                )
                st.plotly_chart(fig_p, use_container_width=True, key=f"pareto_{i}")

        # ── Últimas observaciones de rechazo ──
        st.markdown("---")
        st.markdown("#### Últimas Observaciones de Rechazo")
        st.caption("Observaciones con criticidad R ordenadas por fecha descendente")

        df_rech = df[df['Criticidad'] == 'R'].copy()
        if df_rech.empty:
            st.info("No hay observaciones de rechazo en este archivo.")
        else:
            df_rech['Fecha'] = pd.to_datetime(df_rech['Fecha'], errors='coerce')
            df_rech = df_rech.sort_values('Fecha', ascending=False)
            cols_rech = ['Fecha','Modulo','Vehiculo','SistemaUnidad','Descripcion','CritAmpliado']
            cols_rech = [c for c in cols_rech if c in df_rech.columns]
            tabla_rech = df_rech[cols_rech].head(15).copy()
            tabla_rech['Fecha'] = tabla_rech['Fecha'].dt.strftime('%d/%m/%Y')
            tabla_rech.columns = [{'SistemaUnidad':'Sistema','CritAmpliado':'Criticidad',
                                    'Descripcion':'Descripción'}.get(c,c) for c in tabla_rech.columns]
            st.dataframe(tabla_rech, use_container_width=True, hide_index=True)


# ── TAB 5: EXPLORADOR LIBRE ──
with tab5:
    st.markdown("#### Explorador de Datos")
    st.caption("Seleccioná cualquier variable para los ejes y explorá relaciones entre ellas")

    # Columnas disponibles para explorar
    cols_explorar = {
        'Sistema':      'SistemaUnidad',
        'Criticidad':   'CritAmpliado',
        'Tipo MR':      'MR',
        'Modelo':       'Modelo',
        'Servicio':     'Servicio',
        'Mes':          'Mes',
        'Clasificación':'Clasificacion',
    }
    # Filtrar solo las que existen en el df
    cols_explorar = {k: v for k, v in cols_explorar.items() if v in df.columns}

    ex1, ex2, ex3 = st.columns(3)
    with ex1:
        eje_x = st.selectbox("Eje X (categorías)", list(cols_explorar.keys()), index=0, key="ex_x")
    with ex2:
        eje_color = st.selectbox("Color (segunda variable)", list(cols_explorar.keys()), index=1, key="ex_col")
    with ex3:
        tipo_graf = st.selectbox("Tipo de gráfico", ["Barras agrupadas", "Barras apiladas", "Barras apiladas %"], key="ex_tipo")

    col_x     = cols_explorar[eje_x]
    col_color = cols_explorar[eje_color]

    # Agrupar
    df_exp = df.groupby([col_x, col_color]).size().reset_index(name='Cantidad')

    # Construir kwargs sin barnorm para evitar TypeError
    bar_kwargs = dict(
        x=col_x, y='Cantidad', color=col_color,
        text_auto=True,
        labels={col_x: eje_x, col_color: eje_color, 'Cantidad': 'Observaciones'},
        color_discrete_sequence=['#4fc3f7','#ef5350','#ffa726','#66bb6a','#ab47bc','#26c6da','#81d4fa'],
    )
    if tipo_graf == "Barras apiladas %":
        bar_kwargs['barmode']  = 'stack'
        bar_kwargs['barnorm']  = 'percent'
    elif tipo_graf == "Barras apiladas":
        bar_kwargs['barmode']  = 'stack'
    else:
        bar_kwargs['barmode']  = 'group'
    fig_exp = px.bar(df_exp, **bar_kwargs)
    fig_exp.update_traces(textfont=dict(size=11, color='white'), textposition='inside')
    fig_exp.update_layout(
        **PLOTLY_THEME,
        height=450,
        margin=dict(t=20, b=80, l=10, r=10),
        xaxis=dict(tickangle=-30, **AXIS_STYLE),
        yaxis=AXIS_STYLE,
        legend=dict(orientation='h', y=-0.25),
    )
    st.plotly_chart(fig_exp, use_container_width=True, key="chart_explorador")

    # Tabla de datos del gráfico
    with st.expander("Ver datos del gráfico"):
        pivot_exp = df_exp.pivot_table(
            index=col_x, columns=col_color, values='Cantidad', fill_value=0
        )
        st.dataframe(pivot_exp, use_container_width=True)




# ── TAB 6: INDICADORES ──
with tab6:
    if df_out.empty or 'Desviacion' not in df_out.columns:
        st.info(
            "Este tab requiere un archivo con columnas de referencia (RefMin / RefMax / Relevado). "
            "Los indicadores paramétricos no están disponibles para este formato."
        )
    else:
        # ── Umbrales del semáforo (referencia: ISO 13381-1 / CBM estándar) ──
        UMBRAL_AMARILLO = 10   # % sobre el límite
        UMBRAL_ROJO     = 25   # % sobre el límite

        st.markdown("""
        <div style="background:#1a2235;border-left:4px solid #4fc3f7;padding:12px 16px;
                    border-radius:6px;margin-bottom:16px;font-size:0.85rem;color:#b0bec5">
        <b style="color:#4fc3f7">Base metodológica</b><br>
        Los indicadores se calculan sobre los valores individuales extraídos de cada celda
        (separados por <code>/</code>). Fuentes: <b>UIC 518</b> (dinámica de bogies),
        <b>EN 14363:2016</b> (aceptación de vehículos), <b>ISO 13381-1:2015</b>
        (mantenimiento basado en condición), <b>EN 50126</b> (RAMS ferroviario).
        </div>
        """, unsafe_allow_html=True)

        # ── Calcular indicadores por fila ──
        df_ind = df_out.copy()

        def calc_indicadores(row):
            """
            Para cada fila con desvío, calcula:
            - todos_vals: lista de todos los valores medidos (de R1 y R2 con sub-valores)
            - n_total: cantidad de puntos medidos
            - n_fuera: cuántos están fuera del rango
            - dispersion: max - min de todos los valores
            - pct_desvio: |desvío| / |límite más cercano| × 100
            - promedio: promedio de todos los valores
            """
            todos = parse_multivalor(row.get('Relevado1')) + parse_multivalor(row.get('Relevado2'))
            if not todos:
                return pd.Series({
                    'n_puntos':0,'n_fuera':0,'promedio':None,
                    'dispersion':None,'pct_desvio':None
                })

            ref_min = row['RefMin_num'] if pd.notna(row['RefMin_num']) else None
            ref_max = row['RefMax_num'] if pd.notna(row['RefMax_num']) else None

            n_fuera = sum(
                1 for v in todos
                if (ref_min is not None and v < ref_min) or
                   (ref_max is not None and v > ref_max)
            )

            dispersion = round(max(todos) - min(todos), 2) if len(todos) > 1 else 0
            promedio   = round(sum(todos) / len(todos), 2)

            # % desvío sobre el límite más cercano al desvío máximo
            desv_abs = abs(row['Desviacion'])
            if ref_min is not None and ref_max is not None:
                limite = ref_min if row['Desviacion'] < 0 else ref_max
            elif ref_min is not None:
                limite = ref_min
            elif ref_max is not None:
                limite = ref_max
            else:
                limite = None

            pct = round(desv_abs / abs(limite) * 100, 1) if limite and limite != 0 else None

            return pd.Series({
                'n_puntos':   len(todos),
                'n_fuera':    n_fuera,
                'promedio':   promedio,
                'dispersion': dispersion,
                'pct_desvio': pct,
            })

        ind_cols = df_ind.apply(calc_indicadores, axis=1)
        df_ind = pd.concat([df_ind, ind_cols], axis=1)
        df_ind['RefMin_num'] = pd.to_numeric(df_ind.get('RefMin'), errors='coerce')
        df_ind['RefMax_num'] = pd.to_numeric(df_ind.get('RefMax'), errors='coerce')

        # Semáforo por fila
        def semaforo(pct):
            if pct is None or pd.isna(pct): return '⚪'
            if pct >= UMBRAL_ROJO:           return '🔴'
            if pct >= UMBRAL_AMARILLO:       return '🟡'
            return '🟢'

        df_ind['🚦'] = df_ind['pct_desvio'].apply(semaforo)

        # ── KPIs de indicadores ──
        st.markdown('<div class="section-header">Resumen de indicadores paramétricos</div>',
                    unsafe_allow_html=True)

        ki1, ki2, ki3, ki4 = st.columns(4)
        n_rojo    = (df_ind['🚦'] == '🔴').sum()
        n_amarillo= (df_ind['🚦'] == '🟡').sum()
        n_verde   = (df_ind['🚦'] == '🟢').sum()
        pct_max   = df_ind['pct_desvio'].max()
        disp_max  = df_ind['dispersion'].max()

        ki1.markdown(kpi("🔴 Desvíos Críticos (>25%)",  n_rojo,    "danger"),  unsafe_allow_html=True)
        ki2.markdown(kpi("🟡 Bajo seguimiento (10-25%)", n_amarillo,"warning"), unsafe_allow_html=True)
        ki3.markdown(kpi("🟢 Leve (<10%)",               n_verde,   "success"), unsafe_allow_html=True)
        ki4.markdown(kpi("% Desvío Máximo registrado",   f"{pct_max:.1f}%" if pd.notna(pct_max) else "—", "danger"),
                     unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # ── Ranking de vehículos por índice de condición ──
        st.markdown("#### 🏆 Ranking de Unidades — Índice de Condición")
        st.caption(
            "IC = 100 − promedio de % desvíos de todos los parámetros fuera de rango de la unidad. "
            "Cuanto menor el IC, peor el estado. Metodología derivada de FCI / APTA PR-M-S-016."
        )

        df_ind['_unidad'] = df_ind['Modulo'].apply(
            lambda x: None if (x is None or str(x).strip() in ('','0','nan')) else str(x).strip()
        ).fillna(df_ind['Vehiculo'].astype(str).str.strip())

        ranking = df_ind.groupby('_unidad').agg(
            Observaciones   =('Desviacion',  'count'),
            Desv_Max        =('Desviacion',  lambda x: round(x.abs().max(),  2)),
            Pct_Desvio_Max  =('pct_desvio',  lambda x: round(x.max(),  1) if x.notna().any() else None),
            Pct_Desvio_Prom =('pct_desvio',  lambda x: round(x.mean(), 1) if x.notna().any() else None),
            Dispersion_Max  =('dispersion',  lambda x: round(x.max(),  2) if x.notna().any() else None),
            Criticos        =('Criticidad',  lambda x: (x == 'C').sum()),
        ).reset_index()
        ranking.columns = ['Unidad','Obs.','Desv. Máx.','% Desv. Máx.','% Desv. Prom.','Dispers. Máx.','Críticos']
        # IC = 100 - % desv promedio (mínimo 0)
        ranking['IC'] = (100 - ranking['% Desv. Prom.']).clip(lower=0).round(1)
        ranking['🚦'] = ranking['% Desv. Máx.'].apply(semaforo)
        ranking = ranking.sort_values('IC').reset_index(drop=True)

        st.dataframe(
            ranking[['🚦','Unidad','IC','Obs.','% Desv. Máx.','% Desv. Prom.','Dispers. Máx.','Críticos']],
            use_container_width=True, hide_index=True,
            column_config={
                'IC':            st.column_config.NumberColumn('IC (0-100)', format="%.1f"),
                '% Desv. Máx.': st.column_config.NumberColumn(format="%.1f"),
                '% Desv. Prom.':st.column_config.NumberColumn(format="%.1f"),
                'Dispers. Máx.':st.column_config.NumberColumn(format="%.2f"),
            }
        )

        st.markdown("---")

        # ── Semáforo visual por unidad ──
        st.markdown("#### 🚦 Semáforo por Unidad")
        st.caption(
            f"🟢 < {UMBRAL_AMARILLO}% sobre el límite  |  "
            f"🟡 {UMBRAL_AMARILLO}–{UMBRAL_ROJO}%  |  "
            f"🔴 > {UMBRAL_ROJO}% — Umbrales según ISO 13381-1 / CBM estándar"
        )

        # Pivot semáforo: unidades × sistemas
        sem_pivot = df_ind.groupby(['_unidad','SistemaUnidad'])['pct_desvio'].max().reset_index()
        sem_pivot['color'] = sem_pivot['pct_desvio'].apply(
            lambda x: '#c62828' if pd.notna(x) and x >= UMBRAL_ROJO
                      else '#f57f17' if pd.notna(x) and x >= UMBRAL_AMARILLO
                      else '#2e7d32'
        )
        sem_pivot['label'] = sem_pivot.apply(
            lambda r: f"{r['SistemaUnidad']}<br>{r['pct_desvio']:.1f}%"
                      if pd.notna(r['pct_desvio']) else r['SistemaUnidad'], axis=1
        )

        unidades_unicas  = sorted(sem_pivot['_unidad'].unique())
        sistemas_unicos  = sorted(sem_pivot['SistemaUnidad'].unique())

        fig_sem = go.Figure()
        for sist in sistemas_unicos:
            df_s = sem_pivot[sem_pivot['SistemaUnidad'] == sist]
            fig_sem.add_trace(go.Bar(
                name=sist,
                x=df_s['_unidad'],
                y=[1] * len(df_s),
                marker_color=df_s['color'].tolist(),
                text=df_s['pct_desvio'].apply(
                    lambda x: f"{x:.1f}%" if pd.notna(x) else ""
                ).tolist(),
                textposition='inside',
                textfont=dict(color='white', size=9),
                hovertemplate='<b>%{x}</b><br>Sistema: ' + sist + '<br>% Desvío: %{text}<extra></extra>',
            ))

        fig_sem.update_layout(
            **PLOTLY_THEME, barmode='stack', height=max(300, len(unidades_unicas)*18),
            margin=dict(t=20,b=20,l=10,r=10),
            xaxis=dict(tickangle=-45, **AXIS_STYLE),
            yaxis=dict(visible=False),
            legend=dict(orientation='h', y=1.05),
            showlegend=True
        )
        st.plotly_chart(fig_sem, use_container_width=True, key="chart_semaforo")

        st.markdown("---")

        # ── Gráfico de dispersión por categoría ──
        st.markdown("#### 📊 Dispersión de valores medidos por categoría")
        st.caption(
            "Cada punto es una medición individual. Permite ver la distribución real de los valores "
            "y detectar outliers que el promedio oculta. "
            "Referencia: SPC (Shewhart, 1931) aplicado a mantenimiento ferroviario según EN 50126."
        )

        # Expandir todos los sub-valores en filas individuales
        rows_exp = []
        for _, row in df_out.iterrows():
            todos = parse_multivalor(row.get('Relevado1')) + parse_multivalor(row.get('Relevado2'))
            ref_min = pd.to_numeric(row.get('RefMin'), errors='coerce')
            ref_max = pd.to_numeric(row.get('RefMax'), errors='coerce')
            for v in todos:
                fuera = (pd.notna(ref_min) and v < ref_min) or (pd.notna(ref_max) and v > ref_max)
                rows_exp.append({
                    'Categoría': row['DescAgrupada'],
                    'Valor':     v,
                    'RefMin':    ref_min,
                    'RefMax':    ref_max,
                    'Fuera':     fuera,
                    'Unidad':    row.get('Vehiculo',''),
                })
        df_exp2 = pd.DataFrame(rows_exp)

        if not df_exp2.empty:
            categorias_disp = sorted(df_exp2['Categoría'].dropna().unique().tolist())
            sel_cat = st.multiselect(
                "Filtrar categorías", categorias_disp,
                default=categorias_disp[:6] if len(categorias_disp) > 6 else categorias_disp,
                key="ind_cat"
            )
            df_exp2_f = df_exp2[df_exp2['Categoría'].isin(sel_cat)]

            fig_disp = go.Figure()

            # Puntos dentro del rango
            df_ok = df_exp2_f[~df_exp2_f['Fuera']]
            if not df_ok.empty:
                fig_disp.add_trace(go.Box(
                    x=df_ok['Categoría'], y=df_ok['Valor'],
                    name='Dentro del rango', marker_color='#66bb6a',
                    boxpoints='all', jitter=0.4, pointpos=0,
                    marker=dict(size=5, opacity=0.6),
                    line=dict(color='#66bb6a')
                ))

            # Puntos fuera del rango
            df_out2 = df_exp2_f[df_exp2_f['Fuera']]
            if not df_out2.empty:
                fig_disp.add_trace(go.Box(
                    x=df_out2['Categoría'], y=df_out2['Valor'],
                    name='Fuera del rango', marker_color='#ef5350',
                    boxpoints='all', jitter=0.4, pointpos=0,
                    marker=dict(size=6, opacity=0.8, symbol='circle'),
                    line=dict(color='#ef5350')
                ))

            # Líneas de referencia por categoría
            for cat in sel_cat:
                df_c = df_exp2_f[df_exp2_f['Categoría'] == cat]
                if df_c.empty: continue
                ref_min_c = df_c['RefMin'].dropna().mean()
                ref_max_c = df_c['RefMax'].dropna().mean()
                if pd.notna(ref_min_c):
                    fig_disp.add_shape(type='line',
                        x0=cat, x1=cat, y0=ref_min_c*0.98, y1=ref_min_c,
                        line=dict(color='#ffa726', width=3))
                if pd.notna(ref_max_c):
                    fig_disp.add_shape(type='line',
                        x0=cat, x1=cat, y0=ref_max_c, y1=ref_max_c*1.02,
                        line=dict(color='#ffa726', width=3))

            fig_disp.update_layout(
                **PLOTLY_THEME, height=500,
                margin=dict(t=20,b=100,l=10,r=10),
                xaxis=dict(tickangle=-35, **AXIS_STYLE),
                yaxis=dict(title='Valor medido', **AXIS_STYLE),
                legend=dict(orientation='h', y=1.05),
                boxmode='group'
            )
            st.plotly_chart(fig_disp, use_container_width=True, key="chart_dispersion")

        st.markdown("---")

        # ── Tabla detalle con todos los indicadores ──
        st.markdown("#### 📋 Tabla detalle — Todos los indicadores")
        st.caption(
            "Desvío absoluto: diferencia al límite más cercano · "
            "% Desvío: normalizado sobre el límite (comparación entre parámetros) · "
            "Dispersión: max−min de los puntos medidos (detecta desgaste diferencial) · "
            "N puntos: cantidad de mediciones individuales en la celda"
        )

        fc1, fc2 = st.columns(2)
        with fc1:
            sel_sem = st.multiselect("Semáforo", ['🔴','🟡','🟢'],
                                     default=['🔴','🟡','🟢'], key="ind_sem")
        with fc2:
            sel_sist2 = st.multiselect("Sistema", sorted(df_ind['SistemaUnidad'].dropna().unique()),
                                       default=sorted(df_ind['SistemaUnidad'].dropna().unique()),
                                       key="ind_sist")

        df_tabla = df_ind[
            df_ind['🚦'].isin(sel_sem) &
            df_ind['SistemaUnidad'].isin(sel_sist2)
        ].copy()

        df_tabla['Ref.'] = df_tabla.apply(
            lambda r: f"{r['RefMin_num']}–{r['RefMax_num']}"
                      if pd.notna(r['RefMin_num']) and pd.notna(r['RefMax_num'])
                      else f"≥{r['RefMin_num']}" if pd.notna(r['RefMin_num'])
                      else f"≤{r['RefMax_num']}" if pd.notna(r['RefMax_num'])
                      else "—", axis=1
        )

        cols_show = ['🚦','_unidad','SistemaUnidad','Descripcion','Ref.',
                     'ValorRelevado','Desviacion','pct_desvio','dispersion','n_puntos','Criticidad']
        rename = {
            '_unidad':'Unidad','SistemaUnidad':'Sistema','Descripcion':'Descripción',
            'ValorRelevado':'Valor (peor)','Desviacion':'Desvío abs.',
            'pct_desvio':'% Desvío','dispersion':'Dispersión','n_puntos':'N puntos'
        }
        tabla_final = df_tabla[cols_show].rename(columns=rename).reset_index(drop=True)

        st.markdown(f"**{len(tabla_final)}** observaciones")
        st.dataframe(
            tabla_final,
            use_container_width=True, hide_index=True,
            column_config={
                'Valor (peor)': st.column_config.NumberColumn(format="%.1f"),
                'Desvío abs.':  st.column_config.NumberColumn(format="%.2f"),
                '% Desvío':     st.column_config.NumberColumn(format="%.1f"),
                'Dispersión':   st.column_config.NumberColumn(format="%.2f"),
            }
        )


# ── TAB 7: EXPORTAR ──
with tab7:
    st.markdown("#### Configuración del Informe Word")

    # ── Encabezado del documento ──
    st.markdown("##### Datos del encabezado")
    hdr_col1, hdr_col2 = st.columns(2)
    with hdr_col1:
        hdr_codigo   = st.text_input("Código del informe", placeholder="Ej: SGBV-INF-2025-001")
        hdr_version  = st.text_input("Versión", value="v1.0")
        hdr_linea    = st.text_input("Línea / Contrato", placeholder="Ej: Línea San Martín — 3-LA")
    with hdr_col2:
        logo_file    = st.file_uploader("Banner / Logo (JPG o PNG)", type=["jpg","jpeg","png"],
                                         help="Imagen que aparece en el encabezado de todas las páginas")
        hdr_subger   = st.text_input("Subgerencia", value="Sub Gerencia de Programación y Seguimiento de Mantenimiento de Material Rodante (SPySM)")

    st.markdown("---")

    # ── Secciones (tablas) ──
    st.markdown("##### Tablas a incluir")
    sc1, sc2, sc3 = st.columns(3)
    with sc1:
        inc_crit    = st.checkbox("Distribución por Criticidad",  value=True)
        inc_top15   = st.checkbox("Top 15 Tipos de Falla",        value=True)
    with sc2:
        inc_desvios = st.checkbox("Tabla de Desvíos",             value=True)
        inc_detalle = st.checkbox("Detalle fuera de parámetro",   value=True)
    with sc3:
        inc_concl   = st.checkbox("Conclusiones automáticas",     value=True)

    st.markdown("---")

    # ── Gráficos ──
    st.markdown("##### Gráficos a incluir")
    st.caption("Requiere que `kaleido` esté instalado (`pip install kaleido`). "
               "Si no está disponible los gráficos seleccionados se omiten sin error.")

    gc1, gc2, gc3 = st.columns(3)
    with gc1:
        g_torta    = st.checkbox("Torta — Criticidad",            value=False, key="g_torta")
        g_sistemas = st.checkbox("Barras — Sistemas",             value=False, key="g_sistemas")
        g_mensual  = st.checkbox("Línea mensual + ratio",         value=False, key="g_mensual")
    with gc2:
        g_top15    = st.checkbox("Barras — Top 15 Fallas",        value=False, key="g_top15")
        g_desvios  = st.checkbox("Desvíos — distribución",        value=False, key="g_desvios")
    with gc3:
        g_pareto_fr  = st.checkbox("Pareto — Fuera de rango",     value=False, key="g_pfr")
        g_pareto_aus = st.checkbox("Pareto — Ausencia elementos", value=False, key="g_paus")
        g_pareto_mal = st.checkbox("Pareto — Mal estado",         value=False, key="g_pmal")
        g_explorador = st.checkbox("Explorador libre (actual)",   value=False, key="g_exp")

    st.markdown("---")

    if st.button("📄  Generar y descargar informe Word", type="primary"):
        logo_bytes = logo_file.read() if logo_file else None

        # Construir figuras solo si fueron seleccionadas
        graficos = {}

        def fig_to_png(fig, width=700, height=350):
            """Convierte figura a PNG. Retorna None si falla (kaleido/Chrome no disponible)."""
            try:
                return fig.to_image(format='png', width=width, height=height, engine='kaleido')
            except Exception:
                try:
                    return fig.to_image(format='png', width=width, height=height)
                except Exception:
                    return None

        kaleido_ok = True  # intentamos siempre; fig_to_png maneja errores
        any_graf = any([g_torta, g_sistemas, g_mensual, g_top15, g_desvios,
                        g_pareto_fr, g_pareto_aus, g_pareto_mal, g_explorador])

        # Torta criticidad
        if g_torta:
            crit_counts = df['CritAmpliado'].value_counts()
            fig = go.Figure(go.Pie(
                labels=crit_counts.index, values=crit_counts.values, hole=0.55,
                marker=dict(colors=['#4fc3f7','#ef5350','#ffa726','#66bb6a','#ab47bc']),
            ))
            fig.update_layout(**PLOTLY_THEME, height=350, margin=dict(t=30,b=30,l=30,r=30))
            graficos['torta'] = fig_to_png(fig, 600, 350)

        # Barras sistemas
        if g_sistemas:
            sc = df['SistemaUnidad'].value_counts().reset_index()
            sc.columns = ['Sistema','Cantidad']
            sc['Label'] = sc['Sistema'].map(SISTEMA_LABELS).fillna(sc['Sistema'])
            fig = px.bar(sc, x='Cantidad', y='Label', orientation='h',
                         color='Cantidad', color_continuous_scale='Blues', text='Cantidad')
            fig.update_layout(**PLOTLY_THEME, height=350, margin=dict(t=10,b=10,l=10,r=80),
                              xaxis=dict(range=[0, sc['Cantidad'].max()*1.15], **AXIS_STYLE),
                              yaxis=dict(autorange='reversed', **AXIS_STYLE),
                              coloraxis_showscale=False)
            graficos['sistemas'] = fig_to_png(fig, 700, 350)

        # Línea mensual
        if g_mensual:
            monthly = []
            for m in MONTH_ORDER:
                df_m = df[df['Mes'] == m]
                cnt = len(df_m); vehs = df_m['Vehiculo'].dropna().nunique()
                if cnt > 0:
                    monthly.append({'Mes':m,'Observaciones':cnt,'Vehículos':vehs,
                                    'Ratio':round(cnt/vehs,2) if vehs>0 else 0})
            if monthly:
                mdf = pd.DataFrame(monthly)
                fig = go.Figure()
                fig.add_trace(go.Bar(x=mdf['Mes'], y=mdf['Observaciones'],
                                     name='Obs.', marker_color='rgba(79,195,247,0.4)',
                                     text=mdf['Observaciones'], textposition='outside'))
                fig.add_trace(go.Scatter(x=mdf['Mes'], y=mdf['Ratio'],
                                         name='Ratio Obs/Veh', mode='lines+markers',
                                         line=dict(color='#ffa726',width=2), yaxis='y2'))
                fig.update_layout(**PLOTLY_THEME, height=320, barmode='overlay',
                                  margin=dict(t=20,b=20,l=10,r=60),
                                  xaxis=AXIS_STYLE,
                                  yaxis=dict(title='Obs.', **AXIS_STYLE),
                                  yaxis2=dict(title='Ratio', overlaying='y', side='right',
                                              gridcolor='#1e2a3a', linecolor='#2a3a50'))
                graficos['mensual'] = fig_to_png(fig, 800, 320)

        # Top 15 fallas
        if g_top15:
            tp = df['DescAgrupada'].value_counts().head(15).reset_index()
            tp.columns = ['Falla','Cantidad']
            fig = px.bar(tp, x='Cantidad', y='Falla', orientation='h',
                         color='Cantidad', color_continuous_scale='Blues_r', text='Cantidad')
            fig.update_layout(**PLOTLY_THEME, height=420, margin=dict(t=10,b=10,r=80,l=10),
                              xaxis=dict(range=[0,tp['Cantidad'].max()*1.15], **AXIS_STYLE),
                              yaxis=dict(autorange='reversed', **AXIS_STYLE),
                              coloraxis_showscale=False)
            graficos['top15'] = fig_to_png(fig, 800, 420)

        # Desvíos distribución
        if g_desvios and not df_out.empty and 'Desviacion' in df_out.columns:
            dp = df_out.groupby('DescAgrupada').agg(
                Cantidad=('Desviacion','count'),
                Desv_Max=('Desviacion', lambda x: round(x.abs().max(),2)),
                Desv_Prom=('Desviacion', lambda x: round(x.abs().mean(),2)),
            ).reset_index().sort_values('Cantidad', ascending=False)
            fig = go.Figure()
            fig.add_trace(go.Bar(x=dp['DescAgrupada'], y=dp['Cantidad'],
                                 name='Cantidad', marker_color='#4fc3f7',
                                 text=dp['Cantidad'], textposition='outside', yaxis='y'))
            fig.add_trace(go.Scatter(x=dp['DescAgrupada'], y=dp['Desv_Max'],
                                     name='Desvío Máx.', mode='lines+markers',
                                     line=dict(color='#ef5350',width=2), yaxis='y2'))
            fig.add_trace(go.Scatter(x=dp['DescAgrupada'], y=dp['Desv_Prom'],
                                     name='Desvío Prom.', mode='lines+markers',
                                     line=dict(color='#ffa726',width=2,dash='dot'), yaxis='y2'))
            fig.update_layout(**PLOTLY_THEME, height=400, barmode='group',
                              margin=dict(t=20,b=100,l=10,r=60),
                              xaxis=dict(tickangle=-35, **AXIS_STYLE),
                              yaxis=dict(title='Cantidad', **AXIS_STYLE),
                              yaxis2=dict(title='Desvío', overlaying='y', side='right',
                                          gridcolor='#1e2a3a', linecolor='#2a3a50'))
            graficos['desvios'] = fig_to_png(fig, 900, 400)

        # Paretos
        CLASIF_COLORS = {'fuera de rango':'#ef5350','ausencia de elementos':'#ffa726','mal estado':'#4fc3f7'}
        pareto_cfg = [
            ('pareto_fr',  'fuera de rango',        g_pareto_fr),
            ('pareto_aus', 'ausencia de elementos', g_pareto_aus),
            ('pareto_mal', 'mal estado',            g_pareto_mal),
        ]
        if 'Clasificacion' in df.columns:
            for key, cat, sel in pareto_cfg:
                if not sel: continue
                df_cat = df[df['Clasificacion'].str.strip().str.lower() == cat]
                if df_cat.empty: continue
                conteo = df_cat['SistemaUnidad'].value_counts().reset_index()
                conteo.columns = ['Sistema','Cantidad']
                conteo['Label'] = conteo['Sistema'].map(SISTEMA_LABELS).fillna(conteo['Sistema'])
                conteo['%_acum'] = (conteo['Cantidad'].cumsum()/conteo['Cantidad'].sum()*100).round(1)
                color = CLASIF_COLORS.get(cat, '#4fc3f7')
                fig = go.Figure()
                fig.add_trace(go.Bar(x=conteo['Label'], y=conteo['Cantidad'],
                                     marker_color=color, text=conteo['Cantidad'],
                                     textposition='outside', yaxis='y'))
                fig.add_trace(go.Scatter(x=conteo['Label'], y=conteo['%_acum'],
                                         mode='lines+markers',
                                         line=dict(color='#ffffff',width=2), yaxis='y2'))
                fig.add_hline(y=80, line_dash='dash', line_color='#ffa726',
                              annotation_text='80%', yref='y2')
                fig.update_layout(**PLOTLY_THEME, height=350,
                                  margin=dict(t=30,b=60,l=10,r=50),
                                  xaxis=dict(tickangle=-35, **AXIS_STYLE),
                                  yaxis=dict(title='Cantidad', **AXIS_STYLE),
                                  yaxis2=dict(title='% Acum.', overlaying='y', side='right',
                                              range=[0,110], gridcolor='#1e2a3a', linecolor='#2a3a50'),
                                  title=dict(text=cat.title(), font=dict(color='#e8eaf0')))
                graficos[key] = fig_to_png(fig, 700, 350)

        # Explorador libre (usa la última configuración del tab5)
        if g_explorador:
            try:
                col_x_e     = st.session_state.get('ex_x',     'Sistema')
                col_color_e = st.session_state.get('ex_col',    'Criticidad')
                tipo_e      = st.session_state.get('ex_tipo',   'Barras agrupadas')
                cx  = {'Sistema':'SistemaUnidad','Criticidad':'CritAmpliado','Tipo MR':'MR',
                       'Modelo':'Modelo','Servicio':'Servicio','Mes':'Mes',
                       'Clasificación':'Clasificacion'}.get(col_x_e, 'SistemaUnidad')
                cc  = {'Sistema':'SistemaUnidad','Criticidad':'CritAmpliado','Tipo MR':'MR',
                       'Modelo':'Modelo','Servicio':'Servicio','Mes':'Mes',
                       'Clasificación':'Clasificacion'}.get(col_color_e, 'CritAmpliado')
                df_e = df.groupby([cx, cc]).size().reset_index(name='Cantidad')
                bk = dict(x=cx, y='Cantidad', color=cc, text_auto=True,
                          color_discrete_sequence=['#4fc3f7','#ef5350','#ffa726',
                                                   '#66bb6a','#ab47bc','#26c6da'])
                if tipo_e == "Barras apiladas %":
                    bk['barmode'] = 'stack'; bk['barnorm'] = 'percent'
                elif tipo_e == "Barras apiladas":
                    bk['barmode'] = 'stack'
                else:
                    bk['barmode'] = 'group'
                fig = px.bar(df_e, **bk)
                fig.update_layout(**PLOTLY_THEME, height=400,
                                  margin=dict(t=20,b=80,l=10,r=10),
                                  xaxis=dict(tickangle=-30, **AXIS_STYLE),
                                  yaxis=AXIS_STYLE)
                graficos['explorador'] = fig_to_png(fig, 800, 400)
            except Exception:
                pass

        config = dict(
            codigo=hdr_codigo, version=hdr_version, linea=hdr_linea,
            subger=hdr_subger, logo=logo_bytes,
            inc_crit=inc_crit, inc_top15=inc_top15,
            inc_desvios=inc_desvios, inc_detalle=inc_detalle,
            inc_concl=inc_concl,
            graficos=graficos,
        )
        with st.spinner("Generando documento..."):
            word_bytes = generar_word(df, df_out, config)

        st.success("✅ Informe generado correctamente.")
        linea_safe = (hdr_linea or "Informe").replace(" ","_").replace("/","_")[:30]
        st.download_button(
            label="⬇️  Descargar .docx",
            data=word_bytes,
            file_name=f"Informe_{linea_safe}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
